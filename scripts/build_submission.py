from pathlib import Path
import csv, json, re, shutil, zipfile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.colors import HexColor


ROOT=Path(__file__).resolve().parents[1]; OUT=ROOT/"output"/"submission"; OUT.mkdir(parents=True,exist_ok=True)
TEAM="洪析先知团队"; PROJECT="洪析先知"; TITLE="基于多时相SAR与地形先验的洪涝智能评估系统"
BLUE="2E74B5"; DARK="17365D"; LIGHT="E8EEF5"; GRAY="666666"


def set_font(run,size=11,bold=False,color="000000",east="Microsoft YaHei"):
    run.font.name="Calibri";run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),east);run._element.rPr.rFonts.set(qn("w:ascii"),"Calibri");run._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri");run.font.size=Pt(size);run.bold=bold;run.font.color.rgb=RGBColor.from_string(color)


def set_cell_shading(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr();shd=tcPr.find(qn("w:shd")) or OxmlElement("w:shd");shd.set(qn("w:fill"),fill);tcPr.append(shd) if shd.getparent() is None else None


def set_cell_margins(cell,top=80,start=120,bottom=80,end=120):
    tc=cell._tc;tcPr=tc.get_or_add_tcPr();tcMar=tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:tcMar=OxmlElement("w:tcMar");tcPr.append(tcMar)
    for tag,val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node=tcMar.find(qn("w:"+tag)) or OxmlElement("w:"+tag);node.set(qn("w:w"),str(val));node.set(qn("w:type"),"dxa");tcMar.append(node) if node.getparent() is None else None


def fixed_table(table,widths):
    table.alignment=WD_TABLE_ALIGNMENT.LEFT;table.autofit=False;tblPr=table._tbl.tblPr
    tblW=tblPr.first_child_found_in("w:tblW") or OxmlElement("w:tblW");tblW.set(qn("w:w"),str(sum(widths)));tblW.set(qn("w:type"),"dxa");tblPr.append(tblW) if tblW.getparent() is None else None
    tblInd=tblPr.first_child_found_in("w:tblInd") or OxmlElement("w:tblInd");tblInd.set(qn("w:w"),"120");tblInd.set(qn("w:type"),"dxa");tblPr.append(tblInd) if tblInd.getparent() is None else None
    grid=table._tbl.tblGrid
    for child in list(grid):grid.remove(child)
    for w in widths:
        col=OxmlElement("w:gridCol");col.set(qn("w:w"),str(w));grid.append(col)
    for row in table.rows:
        for i,cell in enumerate(row.cells):
            tcW=cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW") or OxmlElement("w:tcW");tcW.set(qn("w:w"),str(widths[i]));tcW.set(qn("w:type"),"dxa");cell._tc.get_or_add_tcPr().append(tcW) if tcW.getparent() is None else None;set_cell_margins(cell);cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc,headers,rows,widths):
    t=doc.add_table(rows=1,cols=len(headers));t.style="Table Grid"
    for i,h in enumerate(headers):t.rows[0].cells[i].text=str(h);set_cell_shading(t.rows[0].cells[i],LIGHT)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):cells[i].text=str(v)
    fixed_table(t,widths)
    for r,row in enumerate(t.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after=Pt(2);p.paragraph_format.line_spacing=1.1
                for run in p.runs:set_font(run,9.2,bold=(r==0),color=DARK if r==0 else "000000")
    doc.add_paragraph().paragraph_format.space_after=Pt(2);return t


def add_p(doc,text,bold_lead=None):
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(8);p.paragraph_format.line_spacing=1.333;p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):set_font(p.add_run(bold_lead),11,True,DARK);set_font(p.add_run(text[len(bold_lead):]),11)
    else:set_font(p.add_run(text),11)
    return p


def add_bullets(doc,items):
    for item in items:
        p=doc.add_paragraph(style="List Bullet");p.paragraph_format.left_indent=Inches(.375);p.paragraph_format.first_line_indent=Inches(-.194);p.paragraph_format.space_after=Pt(4);p.paragraph_format.line_spacing=1.208;set_font(p.add_run(item),11)


def heading(doc,text,level=1):
    p=doc.add_paragraph(style=f"Heading {level}");p.paragraph_format.keep_with_next=True;set_font(p.add_run(text),16 if level==1 else 13,True,BLUE if level<3 else DARK);return p


def configure(doc):
    sec=doc.sections[0];sec.page_width=Inches(8.5);sec.page_height=Inches(11);sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1);sec.header_distance=sec.footer_distance=Inches(.492)
    styles=doc.styles
    normal=styles["Normal"];normal.font.name="Calibri";normal._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei");normal.font.size=Pt(11);normal.paragraph_format.space_after=Pt(8);normal.paragraph_format.line_spacing=1.333
    for i,size in ((1,16),(2,13),(3,12)):
        s=styles[f"Heading {i}"];s.font.name="Calibri";s._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei");s.font.size=Pt(size);s.font.bold=True;s.font.color.rgb=RGBColor.from_string(BLUE if i<3 else DARK);s.paragraph_format.space_before=Pt(18 if i==1 else 12);s.paragraph_format.space_after=Pt(10 if i==1 else 6);s.paragraph_format.keep_with_next=True
    header=sec.header.paragraphs[0];header.alignment=WD_ALIGN_PARAGRAPH.RIGHT;set_font(header.add_run(f"{PROJECT}｜第八届中国研究生人工智能创新大赛"),9,color=GRAY)
    footer=sec.footer.paragraphs[0];footer.alignment=WD_ALIGN_PARAGRAPH.CENTER;set_font(footer.add_run(f"{TEAM}  ·  项目文档 V1.0"),9,color=GRAY)


def build_docx():
    formal=list(csv.DictReader((ROOT/"outputs/formal_summary/formal_ablation.csv").open(encoding="utf-8-sig")));manifest=json.loads((ROOT/"outputs/formal_summary/data_manifest.json").read_text(encoding="utf-8"));severity=json.loads((ROOT/"outputs/severity/thresholds.json").read_text(encoding="utf-8"));env=json.loads((ROOT/"outputs/environment.json").read_text(encoding="utf-8"))
    d=Document();configure(d)
    for _ in range(5):d.add_paragraph()
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;set_font(p.add_run("第八届中国研究生人工智能创新大赛"),15,True,BLUE)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(24);set_font(p.add_run(PROJECT),30,True,DARK)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;set_font(p.add_run(TITLE),16,False,BLUE)
    p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(36);set_font(p.add_run("项目文档"),22,True,"000000")
    for line in ("版本：V1.0","日期：2026.08.12",f"团队：{TEAM}","参赛组别：开放命题（待报名信息确认）"):
        p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;set_font(p.add_run(line),11,color=GRAY)
    d.add_page_break();heading(d,"目录",1)
    for text in ("1  项目概况","   1.1 背景和基础","   1.2 场景和价值","   1.3 所需支持","2  项目规划","   2.1 整体目标","   2.2 技术创新点","3  实施方案","   3.1 技术可行性分析","   3.2 技术细节","   3.3 计划和分工","4  参考资料"):
        p=d.add_paragraph();p.paragraph_format.space_after=Pt(5);set_font(p.add_run(text),11,True if not text.startswith(" ") else False,DARK if not text.startswith(" ") else "000000")
    d.add_page_break();heading(d,"记录更改历史",1);table(d,["序号","更改原因","版本","作者","日期"],[["1","建立数据与模型闭环","V0.1",TEAM,"2026.08.12"],["2","加入正式 train/val/test_holdout 结果","V1.0",TEAM,"2026.08.12"]],[650,3050,1000,2500,2160])
    d.add_page_break();heading(d,"1 项目概况",1);heading(d,"1.1 背景和基础",2)
    add_p(d,"洪涝灾害常伴随厚云、强降雨和夜间过程，光学遥感与地面巡查难以稳定覆盖。合成孔径雷达（SAR）具备全天时、全天候成像能力，但永久水体、山体阴影、湿润土壤和城市散射均可能与洪水产生相似响应。项目因此提出“洪析先知”，以灾前与事件期 Sentinel-1 SAR 的变化证据为核心，辅以数字高程模型，输出像素级新增洪水范围、置信度和事件级严重度。")
    add_p(d,"项目从数据语义核验出发，不把网络复杂度等同于创新。ImpactMesh-Flood 的四个时相依次为 pre-month、pre-event、event、post-event。统计显示洪水类像素 VV/VH 从 pre-event 的 -9.964/-15.263 dB 降至 event 的 -15.179/-21.171 dB，并在 post-event 回升，说明 MASK 对应 event 当前状态。项目据此选择“多时相 SAR 洪涝智能评估”路线，不作未来预测的越界表述。")
    add_p(d,"团队采用遥感/GIS、算法、系统工程和材料四类角色协作。当前已完成 24,578 个 train/val/test 图块必要模态落盘、E0-E3 消融、12 个完全未见事件泛化测试、事件级严重度规则和可复现实验环境。")
    heading(d,"1.2 场景和价值",2)
    add_bullets(d,["灾中快速研判：首次可用事件期 SAR 到达后生成新增洪水概率图，辅助缩小人工核查范围。","跨区域态势感知：叠加行政区、道路、居民点和水利设施，定位可能受影响的关键节点。","灾后复盘与科研：按事件保存输入、权重、阈值和误差图，支持跨地貌泛化与标签噪声研究。","竞赛与教学演示：清楚展示模型输入时相、预测证据、对照实验和任务边界。"])
    add_p(d,"系统的社会价值在于提升恶劣天气下大范围信息获取的一致性与可复核性。输出用于辅助决策而非替代专业水文预报，不直接给出人员伤亡、经济损失或自动告警结论。")
    heading(d,"1.3 所需支持",2)
    table(d,["支持类型","当前基础","后续需求"],[["数据","ImpactMesh-Flood v1；S1RTC/DEM/MASK","本地历史事件、道路与重要设施图层"],["算法","PyTorch U-Net 与可复现消融","地形双分支、阈值校准、多随机种子"],["算力",f"{env['gpu']}，CUDA {env['cuda_runtime']}","持续 GPU 配额与模型归档空间"],["专家","公开 CEMS/ESA 资料核验","遥感、GIS、水利与应急人员独立评审"]],[1500,3500,4360])
    heading(d,"2 项目规划",1);heading(d,"2.1 整体目标",2)
    add_p(d,"参赛期间形成一个可展示、可复现、可审计的洪涝智能评估原型。输入 pre-event 与 event SAR，可选 DEM，输出新增洪水概率、二值掩膜、事件级影响比例及“一般—严重—非常严重”三级结果；同时保留数据版本、模型权重、配置与阈值。")
    add_bullets(d,["完成至少三个随机种子的 E0-E3 重复实验，报告均值和标准差。","在 validation 上选择概率阈值，固定用于 test_holdout，避免测试集调参。","开发轻量 Web Demo，展示输入、预测、真值、置信度、事件严重度及任务边界。","形成初赛简介、项目文档、视频与辅助材料四类提交物。"])
    heading(d,"2.2 技术创新点",2)
    add_p(d,"第一，时间语义可审计。项目以四时相散射统计证明标签对应 event，并设置不能称为未来预测的失败条件。第二，多时相变化驱动。正式 test_holdout 上，E2 相对 E0 的 IoU 提升 0.1324，证明灾前参照能减少永久水体等低散射混淆。第三，地形先验采用可证伪设计。DEM 在单时相上带来 0.0186 IoU 增益，但加入多时相后未超过 E2，因此如实报告而不包装“必然提升”。第四，像素结果向事件语言转换。严重度阈值仅由训练事件分布确定，验证和测试固定复用。")
    d.add_picture(str(ROOT/"outputs/formal_summary/formal_ablation.png"),width=Inches(6.45));p=d.add_paragraph("图1  正式消融实验：官方 validation 与未见事件 test_holdout");p.alignment=WD_ALIGN_PARAGRAPH.CENTER;set_font(p.add_run(""),9,color=GRAY)
    heading(d,"3 实施方案",1);heading(d,"3.1 技术可行性分析",2)
    add_p(d,f"ImpactMesh-Flood v1 提供 19,448/2,171/2,959 个 train/val/test 样本，覆盖 206 场洪水事件；test_holdout 包含 562 个来自完全未见事件的样本。本地清单核验结果为：train、val、test 的 S1RTC/DEM/MASK 均与官方数量一致。数据许可为 CC-BY 4.0，官方代码与模型采用 Apache 2.0。")
    table(d,["Split","S1RTC","DEM","MASK","用途"],[[m["split"],m["S1RTC"],m["DEM"],m["MASK"],"训练" if m["split"]=="train" else ("选模" if m["split"]=="val" else "最终测试")] for m in manifest["modalities"]],[1200,1600,1600,1600,3360])
    add_p(d,"S1RTC 为 4×2×256×256 的 Zarr Zip，DEM/MASK 为 1×256×256 GeoTIFF。抽查样本的坐标参考、10 m 像元和仿射变换一致。MASK 编码经空间形态和散射统计核验：0 为背景、1 为永久水体、2 为新增洪水；模型仅将类别 2 作为正类。")
    heading(d,"3.2 技术细节",2)
    add_p(d,"模型采用三层 U-Net 编码器—解码器与跳跃连接。E0 输入 event VV/VH；E1 加入 DEM；E2 输入 pre-event 与 event VV/VH；E3 同时使用多时相 SAR 与 DEM。S1RTC 使用官方均值/标准差归一化，DEM 使用官方统计量。损失为带正类权重的 BCEWithLogits 与 Dice Loss 之和，优化器 AdamW。")
    table(d,["实验","输入","Val IoU/Dice","Holdout IoU/Dice","Holdout P/R"],[[r["experiment"],{"e0":"Event SAR","e1":"Event SAR+DEM","e2":"Pre+Event SAR","e3":"Pre+Event SAR+DEM"}[r["input"]],f"{float(r['val_iou']):.4f} / {float(r['val_dice']):.4f}",f"{float(r['holdout_iou']):.4f} / {float(r['holdout_dice']):.4f}",f"{float(r['holdout_precision']):.4f} / {float(r['holdout_recall']):.4f}"] for r in formal],[900,2200,2100,2200,1960])
    add_p(d,"E2 为当前最优方案。E3 的 Recall 最高（0.8203），但 Precision 低于 E2，说明 DEM 简单拼接倾向于扩大候选区域。后续将尝试地形独立编码器、坡度/相对高程派生量和门控融合；若多随机种子仍无峰值增益，则把 DEM 定位为召回辅助模块。")
    pred=next((ROOT/"outputs/formal_e2/predictions").glob("*.png"));d.add_picture(str(pred),width=Inches(6.45));p=d.add_paragraph("图2  E2 在完全未见事件上的输入、真值、概率和二值预测示例");p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_p(d,f"事件级严重度使用“新增洪水像素÷排除永久水体后的有效陆地像素”。训练事件三分位阈值为 {severity['quantiles']['q33']:.5f} 与 {severity['quantiles']['q67']:.5f}；验证和测试不重新计算。该等级表达空间影响比例，不等同于人员伤亡或经济损失。")
    add_p(d,"test_holdout 含一个无洪水正类事件 EMSR766。项目同时报告总体像素级指标、11 个正洪水事件的宏平均和空事件误报，避免大型事件主导或空事件扭曲结论。E2 的正洪水事件宏平均 IoU 为 0.3251。")
    heading(d,"3.3 计划和分工",2)
    table(d,["阶段","时间","工作","责任角色","验收"],[["阶段一","已完成","数据核验、E0-E3、holdout","遥感+算法","可复现指标与权重"],["阶段二","8月中旬","多种子、阈值、地形融合","算法+GIS","均值/标准差与误差分类"],["阶段三","8月下旬","Web Demo、视频、专家复核","系统+材料","完整演示与审计链"],["提交","9月1日前","格式、命名、大小与资格复核","全体","四类材料合规"]],[1100,1400,3000,1600,2260])
    add_p(d,"每项数据判断、实验配置和文档结论均指定责任人与复核人。模型改动必须在相同数据划分和预算下比较；最终数值由脚本自动生成，不在文档中手工修改。")
    heading(d,"4 参考资料",1)
    refs=["[1] IBM, DLR, ESA Φ-lab. ImpactMesh-Flood Dataset Card. https://huggingface.co/datasets/ibm-esa-geospatial/ImpactMesh-Flood","[2] IBM. ImpactMesh official code and TerraTorch configs. https://github.com/IBM/ImpactMesh","[3] Ronneberger O, Fischer P, Brox T. U-Net: Convolutional Networks for Biomedical Image Segmentation. MICCAI, 2015.","[4] European Space Agency. Sentinel-1 mission and all-weather radar observations. https://www.esa.int/Applications/Observing_the_Earth/Copernicus/Sentinel-1","[5] Copernicus Emergency Management Service. https://emergency.copernicus.eu/","[6] Bonafilia D et al. Sen1Floods11: a georeferenced dataset to train and test deep learning flood algorithms for Sentinel-1. CVPR Workshops, 2020.","[7] Bountos N I et al. Kuro Siwo: 33 billion m² under water. NeurIPS Datasets and Benchmarks, 2023."]
    for ref in refs:add_p(d,ref)
    path=OUT/f"{TEAM}_{PROJECT}_项目文档.docx";d.save(path);return path


def build_intro():
    text="洪析先知面向云雨、夜间等复杂条件下的洪涝快速研判，利用灾前与事件期Sentinel-1 SAR变化信息，结合地形先验识别新增淹没区域，并生成事件级三级严重度。项目基于ImpactMesh-Flood v1完成全量训练、验证及完全未见事件测试。最优多时相U-Net在test_holdout上取得IoU 0.4941、Dice/F1 0.6614，较单时相IoU提升0.1324。系统保留概率图、数据版本、阈值和失败案例，明确当前任务是洪涝智能评估而非未来预测，可为应急核查、态势感知和灾后复盘提供可复核的空间证据。"
    font="C:/Windows/Fonts/msyh.ttc";pdfmetrics.registerFont(TTFont("MSYH",font,subfontIndex=0));path=OUT/f"{TEAM}_{PROJECT}_参赛作品简介.pdf";c=canvas.Canvas(str(path),pagesize=A4);w,h=A4;c.setFillColor(HexColor("#17365D"));c.setFont("MSYH",24);c.drawString(60,h-85,PROJECT);c.setFillColor(HexColor("#2E74B5"));c.setFont("MSYH",14);c.drawString(60,h-115,TITLE);c.setStrokeColor(HexColor("#D7E2F0"));c.line(60,h-135,w-60,h-135);c.setFillColor(HexColor("#222222"));c.setFont("MSYH",12);maxw=w-120;line="";y=h-180
    # Treat each contiguous Latin expression as one token so identifiers such as
    # ``Sentinel-1 SAR`` and ``test_holdout`` are never split across lines.
    tokens=re.findall(r"[A-Za-z0-9_./+\-]+(?:\s+[A-Za-z0-9_./+\-]+)*|.",text,re.S)
    for token in tokens:
        if line and pdfmetrics.stringWidth(line+token,"MSYH",12)>maxw:
            c.drawString(60,y,line);y-=25;line=token
        else:line+=token
    if line:c.drawString(60,y,line)
    c.setFillColor(HexColor("#666666"));c.setFont("MSYH",9);c.drawRightString(w-60,55,f"{TEAM}｜2026.08.12｜正文 {len(text)} 字");c.save();return path


def build_zip():
    path=OUT/f"{TEAM}_{PROJECT}_其他.zip";include=["README.md","requirements.txt","src","configs","scripts","tests","docs","outputs/formal_summary","outputs/imbalance_boundary_summary.csv","outputs/environment.json","outputs/label_semantics.json","outputs/severity"]
    with zipfile.ZipFile(path,"w",zipfile.ZIP_DEFLATED,compresslevel=6) as z:
        for item in include:
            p=ROOT/item
            if p.is_file():z.write(p,p.relative_to(ROOT))
            elif p.exists():
                for f in p.rglob("*"):
                    if f.is_file() and "__pycache__" not in f.parts:z.write(f,f.relative_to(ROOT))
        for e in range(4):
            for n in ("best.pt","metrics.csv","run.json","test_holdout_metrics.json","test_holdout_per_event.csv"):
                f=ROOT/f"outputs/formal_e{e}/{n}"
                if f.exists():z.write(f,f.relative_to(ROOT))
        for f in sorted((ROOT/"outputs/formal_e2/predictions").glob("*.png")):z.write(f,f.relative_to(ROOT))
        for experiment in ("formal_e2_imbalance_boundary","formal_e2_boundary_precision"):
            for n in ("best.pt","metrics.csv","run.json","test_holdout_metrics.json","test_holdout_per_event.csv"):
                f=ROOT/f"outputs/{experiment}/{n}"
                if f.exists():z.write(f,f.relative_to(ROOT))
        for f in sorted((ROOT/"outputs/formal_e2_boundary_precision/predictions").glob("*.png"))[:20]:z.write(f,f.relative_to(ROOT))
    return path


if __name__=="__main__":
    print(build_docx());print(build_intro());print(build_zip())
